# aiubetesis/Aplicaciones/document_service.py
"""
Servicio de exportación de documentos.
Genera archivos Word (.docx) y PDF a partir del diagnostico_json
almacenado en ia.solucion_generada.

Estructura esperada en diagnostico_json (DiagnosticoIADTO):
  {
    "materia_detectada": str,
    "tema_critico_detectado": str,
    "nivel_riesgo": str,
    "justificacion_pedagogica": [str],
    "plan_refuerzo_eva": [
        {
            "tipo": str,
            "titulo": str,
            "objetivo": str,
            "actividades_pasos": [str],
            "orientaciones_metodologicas": [str],
            "bibliografia": [str],
            "rubrica_evaluacion": [
                {"criterio": str, "descripcion": str, "puntos_maximos": int}
            ]
        }
    ],
    "conclusion": str,
    "requiere_intervencion_docente": bool
  }
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ExportDocumentService:
    def __init__(self, solucion):
        """
        :param solucion: instancia de SolucionGenerada
        """
        self.solucion = solucion
        self.data = solucion.diagnostico_json

    # ----------------------------------------------------------
    # WORD
    # ----------------------------------------------------------
    def generar_word(self) -> io.BytesIO:
        doc = Document()
        self._configurar_estilos(doc)
        self._agregar_encabezado(doc)
        self._agregar_diagnostico(doc)
        self._agregar_plan_refuerzo(doc)
        self._agregar_conclusion(doc)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _configurar_estilos(self, doc):
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

    def _agregar_encabezado(self, doc):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("UNIVERSIDAD BOLIVARIANA DEL ECUADOR")
        run.bold = True
        run.font.size = Pt(13)

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("PLAN DE REFUERZO ACADÉMICO — BAJO RENDIMIENTO ESTUDIANTIL")
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph()

        # Tabla de datos generales
        table = doc.add_table(rows=5, cols=2)
        table.style = "Table Grid"
        datos = [
            ("Carrera",          self.solucion.carrera_nombre),
            ("Asignatura",       self.solucion.asignatura_nombre),
            ("Código",           self.solucion.asignatura_codigo),
            ("Docente",          self.solucion.docente_nombre),
            ("Período",          self.solucion.periodo_codigo),
        ]
        for i, (etiq, val) in enumerate(datos):
            row = table.rows[i]
            row.cells[0].text = etiq
            row.cells[1].text = val
            for run in row.cells[0].paragraphs[0].runs:
                run.bold = True

        doc.add_paragraph()

        # Métricas de rendimiento
        p = doc.add_paragraph()
        p.add_run(
            f"Estudiantes analizados: {self.solucion.total_estudiantes}   |   "
            f"Con bajo rendimiento: {self.solucion.cantidad_estudiantes_recuperacion}   |   "
            f"Nivel de riesgo: {self.data.get('nivel_riesgo', 'N/A').upper()}"
        ).italic = True

        doc.add_paragraph()

    def _agregar_diagnostico(self, doc):
        h = doc.add_heading("DIAGNÓSTICO PSICOPEDAGÓGICO", level=1)

        p = doc.add_paragraph()
        p.add_run("Tema crítico detectado: ").bold = True
        p.add_run(self.data.get("tema_critico_detectado", ""))

        p = doc.add_paragraph()
        p.add_run("Componente principal: ").bold = True
        p.add_run(self.data.get("componente_principal", ""))

        doc.add_paragraph().add_run("Justificación pedagógica:").bold = True
        for item in self.data.get("justificacion_pedagogica", []):
            doc.add_paragraph(f"• {item}", style="List Bullet")

        doc.add_paragraph()

    def _agregar_plan_refuerzo(self, doc):
        doc.add_heading("PLAN DE REFUERZO EVA", level=1)
        tareas = self.data.get("plan_refuerzo_eva", [])

        for idx, tarea in enumerate(tareas):
            doc.add_heading(
                f"Tarea Práctica {idx + 1}: {tarea.get('titulo', '')}",
                level=2
            )

            # Tipo y objetivo
            p = doc.add_paragraph()
            p.add_run("Tipo: ").bold = True
            p.add_run(tarea.get("tipo", ""))

            p = doc.add_paragraph()
            p.add_run("Objetivo: ").bold = True
            p.add_run(tarea.get("objetivo", ""))

            # Pasos
            doc.add_paragraph().add_run("Actividades:").bold = True
            for i, paso in enumerate(tarea.get("actividades_pasos", []), 1):
                doc.add_paragraph(f"Situación {i}: {paso}", style="List Number")

            # Orientaciones
            doc.add_paragraph().add_run("Orientaciones metodológicas:").bold = True
            for ori in tarea.get("orientaciones_metodologicas", []):
                doc.add_paragraph(f"• {ori}", style="List Bullet")

            # Bibliografía
            doc.add_paragraph().add_run("Bibliografía:").bold = True
            for bib in tarea.get("bibliografia", []):
                doc.add_paragraph(f"• {bib}", style="List Bullet")

            # Rúbrica
            doc.add_paragraph().add_run("Rúbrica de Evaluación:").bold = True
            rubricas = tarea.get("rubrica_evaluacion", [])
            if rubricas:
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                hdrs = table.rows[0].cells
                hdrs[0].text = "Criterio"
                hdrs[1].text = "Descripción"
                hdrs[2].text = "Puntos"
                for cell in hdrs:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

                for r in rubricas:
                    row = table.add_row().cells
                    row[0].text = r.get("criterio", "")
                    row[1].text = r.get("descripcion", "")
                    row[2].text = str(r.get("puntos_maximos", 0))

            if idx < len(tareas) - 1:
                doc.add_page_break()
            else:
                doc.add_paragraph()

    def _agregar_conclusion(self, doc):
        doc.add_heading("CONCLUSIÓN", level=1)
        doc.add_paragraph(self.data.get("conclusion", ""))

        intervencion = self.data.get("requiere_intervencion_docente", False)
        p = doc.add_paragraph()
        p.add_run("Requiere intervención docente: ").bold = True
        p.add_run("SÍ" if intervencion else "NO")

    # ----------------------------------------------------------
    # PDF (generado desde el HTML del Word via reportlab)
    # ----------------------------------------------------------
    def generar_pdf(self) -> io.BytesIO:
        """
        Genera un PDF usando reportlab.
        Si reportlab no está disponible se devuelve el Word convertido a bytes
        con cabecera de error informativa.
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Table,
                TableStyle, PageBreak,
            )
            from reportlab.lib import colors

            buffer = io.BytesIO()
            doc_pdf = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=2 * cm,
                leftMargin=2 * cm,
                topMargin=2 * cm,
                bottomMargin=2 * cm,
            )

            styles = getSampleStyleSheet()
            titulo_style = ParagraphStyle(
                "Titulo",
                parent=styles["Heading1"],
                fontSize=13,
                spaceAfter=6,
                textColor=colors.HexColor("#1a3c6e"),
            )
            subtitulo_style = ParagraphStyle(
                "Subtitulo",
                parent=styles["Heading2"],
                fontSize=11,
                spaceAfter=4,
                textColor=colors.HexColor("#2e6da4"),
            )
            normal_style = styles["Normal"]
            normal_style.fontSize = 10

            story = []

            # Encabezado
            story.append(Paragraph("UNIVERSIDAD BOLIVARIANA DEL ECUADOR", titulo_style))
            story.append(Paragraph(
                "PLAN DE REFUERZO ACADÉMICO — BAJO RENDIMIENTO ESTUDIANTIL",
                subtitulo_style
            ))
            story.append(Spacer(1, 0.3 * cm))

            # Datos generales
            datos_tabla = [
                ["Carrera",    self.solucion.carrera_nombre],
                ["Asignatura", self.solucion.asignatura_nombre],
                ["Código",     self.solucion.asignatura_codigo],
                ["Docente",    self.solucion.docente_nombre],
                ["Período",    self.solucion.periodo_codigo],
            ]
            t = Table(datos_tabla, colWidths=[4 * cm, 13 * cm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#d0e4f7")),
                ("FONTNAME",   (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTSIZE",   (0, 0), (-1, -1), 9),
                ("GRID",       (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
            ]))
            story.append(t)
            story.append(Spacer(1, 0.5 * cm))

            # Métricas
            story.append(Paragraph(
                f"<b>Estudiantes analizados:</b> {self.solucion.total_estudiantes} &nbsp;|&nbsp; "
                f"<b>Con bajo rendimiento:</b> {self.solucion.cantidad_estudiantes_recuperacion} &nbsp;|&nbsp; "
                f"<b>Nivel de riesgo:</b> {self.data.get('nivel_riesgo','N/A').upper()}",
                normal_style
            ))
            story.append(Spacer(1, 0.5 * cm))

            # Diagnóstico
            story.append(Paragraph("DIAGNÓSTICO PSICOPEDAGÓGICO", titulo_style))
            story.append(Paragraph(
                f"<b>Tema crítico:</b> {self.data.get('tema_critico_detectado','')}",
                normal_style
            ))
            story.append(Paragraph(
                f"<b>Componente principal:</b> {self.data.get('componente_principal','')}",
                normal_style
            ))
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph("<b>Justificación pedagógica:</b>", normal_style))
            for item in self.data.get("justificacion_pedagogica", []):
                story.append(Paragraph(f"• {item}", normal_style))

            story.append(Spacer(1, 0.5 * cm))
            story.append(Paragraph("PLAN DE REFUERZO EVA", titulo_style))

            for idx, tarea in enumerate(self.data.get("plan_refuerzo_eva", [])):
                story.append(Paragraph(
                    f"Tarea Práctica {idx+1}: {tarea.get('titulo','')}",
                    subtitulo_style
                ))
                story.append(Paragraph(
                    f"<b>Tipo:</b> {tarea.get('tipo','')} &nbsp; "
                    f"<b>Objetivo:</b> {tarea.get('objetivo','')}",
                    normal_style
                ))
                story.append(Spacer(1, 0.2 * cm))

                for i, paso in enumerate(tarea.get("actividades_pasos", []), 1):
                    story.append(Paragraph(f"{i}. {paso}", normal_style))

                story.append(Spacer(1, 0.2 * cm))
                story.append(Paragraph("<b>Orientaciones:</b>", normal_style))
                for ori in tarea.get("orientaciones_metodologicas", []):
                    story.append(Paragraph(f"• {ori}", normal_style))

                story.append(Paragraph("<b>Bibliografía:</b>", normal_style))
                for bib in tarea.get("bibliografia", []):
                    story.append(Paragraph(f"• {bib}", normal_style))

                # Rúbrica
                rubricas = tarea.get("rubrica_evaluacion", [])
                if rubricas:
                    story.append(Paragraph("<b>Rúbrica de Evaluación:</b>", normal_style))
                    rub_data = [["Criterio", "Descripción", "Puntos"]]
                    for r in rubricas:
                        rub_data.append([
                            r.get("criterio", ""),
                            r.get("descripcion", ""),
                            str(r.get("puntos_maximos", 0)),
                        ])
                    rub_t = Table(rub_data, colWidths=[4 * cm, 11 * cm, 2 * cm])
                    rub_t.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2e6da4")),
                        ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
                        ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE",   (0, 0), (-1, -1), 8),
                        ("GRID",       (0, 0), (-1, -1), 0.5, colors.grey),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                         [colors.white, colors.HexColor("#f0f6fc")]),
                    ]))
                    story.append(rub_t)

                story.append(Spacer(1, 0.3 * cm))
                if idx < len(self.data.get("plan_refuerzo_eva", [])) - 1:
                    story.append(PageBreak())

            # Conclusión
            story.append(Paragraph("CONCLUSIÓN", titulo_style))
            story.append(Paragraph(self.data.get("conclusion", ""), normal_style))
            intervencion = self.data.get("requiere_intervencion_docente", False)
            story.append(Paragraph(
                f"<b>Requiere intervención docente:</b> {'SÍ' if intervencion else 'NO'}",
                normal_style
            ))

            doc_pdf.build(story)
            buffer.seek(0)
            return buffer

        except ImportError:
            # Fallback: devuelve el Word en caso de que reportlab no esté instalado
            return self.generar_word()